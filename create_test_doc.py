"""
Create a test multilingual document for the Doc_LLM application
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Add title
title = doc.add_heading('South Indian Languages Test Document', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add introduction
intro = doc.add_paragraph()
intro.add_run('This document contains text in multiple South Indian languages to test the multilingual document processing capabilities.\n\n').bold = True

# Malayalam section
doc.add_heading('Malayalam (മലയാളം)', level=1)
malayalam_text = """
കേരളം ഭാരതത്തിന്റെ തെക്ക് പടിഞ്ഞാറൻ തീരത്തുള്ള ഒരു സംസ്ഥാനമാണ്. 
കേരളത്തിന്റെ തലസ്ഥാനം തിരുവനന്തപുരം ആണ്. 
കേരളത്തെ "ദൈവത്തിന്റെ സ്വന്തം നാട്" എന്ന് വിളിക്കുന്നു.
കേരളത്തിൽ മനോഹരമായ കടൽത്തീരങ്ങളും കായൽ പ്രദേശങ്ങളും പർവ്വത പ്രദേശങ്ങളും ഉണ്ട്.
കേരള സംസ്കാരം വളരെ സമ്പന്നവും വൈവിധ്യപൂർണ്ണവുമാണ്.
"""
doc.add_paragraph(malayalam_text)

# Tamil section
doc.add_heading('Tamil (தமிழ்)', level=1)
tamil_text = """
தமிழ்நாடு இந்தியாவின் தென்கிழக்கு பகுதியில் அமைந்துள்ள ஒரு மாநிலம். 
தலைநகரம் சென்னை ஆகும். 
தமிழ் மொழி உலகின் பழமையான மொழிகளில் ஒன்றாகும்.
தமிழ்நாட்டில் பல பண்டைய கோவில்கள் உள்ளன.
தமிழ் இலக்கியம் மிகவும் செழுமையானது.
"""
doc.add_paragraph(tamil_text)

# Telugu section
doc.add_heading('Telugu (తెలుగు)', level=1)
telugu_text = """
తెలుగు భారతదేశంలో అత్యధికంగా మాట్లాడే భాషలలో ఒకటి.
ఆంధ్రప్రదేశ్ మరియు తెలంగాణ రాష్ట్రాల అధికార భాష తెలుగు.
తెలుగు సాహిత్యం చాలా గొప్పది.
హైదరాబాద్ తెలంగాణ రాష్ట్ర రాజధాని.
విజయవాడ ఆంధ్రప్రదేశ్ లో ఒక ముఖ్యమైన నగరం.
"""
doc.add_paragraph(telugu_text)

# Kannada section
doc.add_heading('Kannada (ಕನ್ನಡ)', level=1)
kannada_text = """
ಕರ್ನಾಟಕ ದಕ್ಷಿಣ ಭಾರತದಲ್ಲಿರುವ ಒಂದು ರಾಜ್ಯ.
ಬೆಂಗಳೂರು ಕರ್ನಾಟಕದ ರಾಜಧಾನಿ.
ಕನ್ನಡ ಭಾಷೆ ಬಹಳ ಪುರಾತನವಾದದ್ದು.
ಕರ್ನಾಟಕದಲ್ಲಿ ಅನೇಕ ಐತಿಹಾಸಿಕ ಸ್ಥಳಗಳಿವೆ.
ಕನ್ನಡ ಸಾಹಿತ್ಯ ಬಹಳ ಶ್ರೀಮಂತವಾಗಿದೆ.
"""
doc.add_paragraph(kannada_text)

# English section
doc.add_heading('About This Document', level=1)
english_text = """
This test document demonstrates the multilingual capabilities of the South Indian Document QA Chatbot.

The system can:
1. Automatically detect the language of uploaded documents
2. Extract text from various formats (PDF, DOCX, TXT, Images)
3. Process and understand content in Malayalam, Tamil, Telugu, Kannada, and Tulu
4. Answer questions in the same language as the document
5. Use advanced RAG (Retrieval-Augmented Generation) for accurate responses

Technologies used:
- Streamlit for the user interface
- Sentence Transformers for multilingual embeddings
- Qdrant for vector storage
- Groq API for language model inference
- Tesseract OCR for image text extraction
"""
doc.add_paragraph(english_text)

# Add footer
doc.add_paragraph('\n')
footer = doc.add_paragraph('Generated for testing multilingual document processing')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save the document
doc.save('test_multilingual_doc.docx')
print("✅ Created test_multilingual_doc.docx successfully!")
print("\nYou can now upload this file to the chatbot to test:")
print("1. Language detection")
print("2. Text extraction from DOCX")
print("3. Multilingual question answering")
